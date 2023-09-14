import os
from flask import Flask, request, redirect, url_for, render_template, jsonify

from PDFparser import PdfParser
from Translator import googleTranslator
from fileManager import *
import config
from docx import Document

app = Flask(__name__)
app.jinja_env.filters['zip'] = zip

upload_path = './upload/'

historyManager = LogHistory()

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return 'No file part in the request'
    
    file = request.files['file']

    if file.filename == '' :
        return 'No selected file'
    
    if file:
        checkDirExist(config.data_path)

        # Create a module to parse the PDF file
        parser = PdfParser(file)
        parser.parseWithOutBound()

        # Crate a module to translate the text
        trans = googleTranslator()
        trans.translate('zh-tw')

        # return redirect(url_for('test'))
        return redirect(url_for('result'))
   
@app.route('/result')
def result():
    transManager = fileManager(os.path.join(config.data_path, 'translated_text.json'))
    parsedManager = fileManager(os.path.join(config.data_path, 'parsed_text.json'))

    translated_data = transManager.readTheFile()
    parsed_data = parsedManager.readTheFile()
    return render_template('result.html',text=translated_data, origin=parsed_data)

@app.route('/delete', methods=['POST'])
def delete() :
    print("Let start removing the useless text")
    indices = request.form.getlist('indices[]')
    indices = [tuple(int(num) for num in index.split('-') )for index in indices]
    
    translated_text = fileManager(os.path.join(config.data_path, 'translated_text.json')) 
    parsed_text = fileManager(os.path.join(config.data_path, 'parsed_text.json'))

    # Start to remove the index that is useless for user
    try:
        translated_text.deleteData(indices)
        parsed_text.deleteData(indices)

        historyManager.push(Log(indices,(-1,-1),'delete'))
        return jsonify(success=True)
    except Exception as e:
        return jsonify(success=False, error=e)

@app.route('/translate', methods=['POST'])
def translate():
    print("Let start to translate a new merged text")
    text = request.form['text']
    indices = request.form.getlist('indices[]')
    indices = [tuple(int(num) for num in index.split('-') )for index in indices]

    indices = sorted(indices, key=lambda x : (x[0],x[1]))
    # Get the origin data 
    transTextManager = fileManager(os.path.join(config.data_path, 'translated_text.json'))
    parsedTextManager = fileManager(os.path.join(config.data_path, 'parsed_text.json'))

    # Start to delete the content in the origin dictionary and put the new translated text in the corresponding index
    try :
        translator = googleTranslator(text=text)
        translated_merged_text = translator.translate_merged('zh-tw')
        
        position = transTextManager.getData(indices[0][0],indices[0][1],'position')
        translated_data = {
            "page_index" : indices[0][0],
            "element_index" : indices[0][1],
            "text" : translated_merged_text,
            "position" : position
        }
        parsed_data = {
            "page_index" : indices[0][0],
            "element_index" : indices[0][1],
            "text" : text,
            "position" : position
        }

        transTextManager.deleteData(indices)
        parsedTextManager.deleteData(indices)

        transTextManager.addData([translated_data])
        parsedTextManager.addData([parsed_data])

        historyManager.push(Log(indices,indices[0],'merge'))
        return jsonify(success=True, text=translated_merged_text, indices = indices)
    except Exception as e :
        return jsonify(success=False, error=e)

@app.route('/download', methods=['GET'])
def download():
    transTextManager = fileManager(os.path.join(config.data_path, 'translated_text.json'))
    parsedTextManager = fileManager(os.path.join(config.data_path, 'parsed_text.json'))
    translated_text = transTextManager.readTheFile()
    parsed_text = parsedTextManager.readTheFile()

    try:
        doc = Document()
        doc.add_heading("The Translated text from PDFTranslator", level=1)
        for (trans_value, parse_value)in zip(translated_text,parsed_text):
            parse_value = parse_value['text'].replace('\f',' ')
            trans_value = trans_value['text'].replace('\f',' ')

            doc.add_paragraph("Origin / Translated Text", style='List Number')
            doc.add_paragraph(parse_value, style='Normal')
            doc.add_paragraph(trans_value, style='Normal')
        

        checkDirExist(config.download_path)
        doc.save('./download/output.docx')
        return jsonify(success=True)
    except Exception as e:
        print(e)
        return jsonify(success=False, error = e)

@app.route('/back', methods=['GET'])
def back() :
    try :
        # print(len(historyManager.getHistory()))

        # Get the backup file manager
        backupParsedManager = fileManager(os.path.join(config.data_path, 'backup_parsedText.json'))
        backupTransManager = fileManager(os.path.join(config.data_path,'backup_transText.json'))

        # Get the current file manager
        transTextManager = fileManager(os.path.join(config.data_path, 'translated_text.json'))
        parsedTextManager = fileManager(os.path.join(config.data_path, 'parsed_text.json'))

        # get the previous log file
        log = historyManager.back()

        if log.action == 'none' :
            return jsonify(success=True, undo='nothing')
        # print("I get {} , {} and {}".format(log.indices, log.mergedIndex,log.action))
        backupParsedList = []
        backupTransList = []

        # retrieve the previous data to the list
        for (pageIndex, elementIndex) in log.indices :
            backupParsedList.append(backupParsedManager.getData(pageIndex,elementIndex,'data'))
            backupTransList.append(backupTransManager.getData(pageIndex,elementIndex,'data'))
        # revert the data which is merged before
        if log.action == 'merge' :
            print("remove the merge data {}".format(log.mergedIndex))
            transTextManager.deleteData([log.mergedIndex])
            parsedTextManager.deleteData([log.mergedIndex])

        # store the previus data into the file
        transTextManager.addData(backupTransList)
        parsedTextManager.addData(backupParsedList)
        return jsonify(success=True)
    except Exception as e:
        return jsonify(success=False, error=e)

if __name__ == '__main__':
    app.run(host='localhost',port=4000,debug=True)